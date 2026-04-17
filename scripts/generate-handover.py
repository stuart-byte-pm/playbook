"""
Generate the Playbook Advisory Group client handover .docx document.
Uses python-docx. Run from the project root.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ---------------------------------------------------------------------------
# Page setup - A4, 2cm margins
# ---------------------------------------------------------------------------
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)

# ---------------------------------------------------------------------------
# Colours
# ---------------------------------------------------------------------------
CHARCOAL = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
style_normal = doc.styles["Normal"]
style_normal.font.name = "Arial"
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = CHARCOAL
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for level, size, color in [
    ("Heading 1", 20, CHARCOAL),
    ("Heading 2", 14, CHARCOAL),
    ("Heading 3", 11, DARK_GREY),
]:
    s = doc.styles[level]
    s.font.name = "Arial"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(24 if level == "Heading 1" else 18 if level == "Heading 2" else 12)
    s.paragraph_format.space_after = Pt(10 if level == "Heading 1" else 6 if level == "Heading 2" else 4)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def add_para(text, style="Normal", bold=False, color=None, size=None,
             alignment=None, space_before=None, space_after=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if alignment is not None:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = "Arial"
        rb.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        '<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls("w"), color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = CHARCOAL
        set_cell_shading(cell, "C8F135")
    # Body
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            set_cell_shading(cell, "F7F7F7" if r_idx % 2 == 0 else "FFFFFF")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_placeholder(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    run.bold = True
    rPr = run._element.get_or_add_rPr()
    highlight = parse_xml(
        '<w:highlight {} w:val="yellow"/>'.format(nsdecls("w"))
    )
    rPr.append(highlight)
    return p


# =========================================================================
# COVER PAGE
# =========================================================================
add_para(
    "CONFIDENTIAL",
    alignment=WD_ALIGN_PARAGRAPH.RIGHT, color=GREY, size=Pt(9), space_before=Pt(0),
)
add_para("", space_before=Pt(180))
add_para(
    "Playbook Advisory Group",
    style="Heading 1", alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=Pt(0),
)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Website build \u2014 client handover")
run_sub.font.name = "Arial"
run_sub.font.size = Pt(16)
run_sub.font.color.rgb = GREY
p_sub.paragraph_format.space_after = Pt(120)

add_para("Prepared by Byte  |  byte-pm.com", color=GREY, size=Pt(11))
add_para("April 2026", color=GREY, size=Pt(11), space_after=Pt(0))
doc.add_page_break()


# =========================================================================
# 1. WELCOME & OVERVIEW
# =========================================================================
doc.add_heading("Welcome & overview", level=1)
add_para(
    "This document is a complete handover pack for the Playbook Advisory Group "
    "marketing website at playbook-group.co.uk. It covers everything you need to "
    "know about what was built, how it works, where things are hosted, and how to "
    "manage content going forward."
)
add_para(
    "The site positions Playbook Advisory Group as a mature, boutique advisory "
    "practice from day one \u2014 a senior-led, sponsor-side firm specialising in "
    "capital programme governance. It is designed to serve sponsors, senior leaders, "
    "and programme directors across public and private sectors who need clarity, "
    "control, and confidence over complex capital programmes."
)
add_para(
    "The website was designed and built by Byte (byte-pm.com). If you have any "
    "questions about anything in this document, please get in touch \u2014 our "
    "contact details are at the end."
)


# =========================================================================
# 2. PROJECT SCOPE SUMMARY
# =========================================================================
doc.add_heading("Project scope summary", level=1)
add_para("The following deliverables were completed as part of the Phase 1 build:")

delivered_items = [
    "Fully responsive homepage with hero video, belief section, \u201cWhere we sit\u201d "
    "manifesto, engagement cards, services accordion, sector cards, insights feed, "
    "diagnostic CTA, contact form, and footer",
    "Three sector pages (Public, Private, Infrastructure) with 12 sub-sector "
    "sections, alternating two-column layouts, and shared diagnostic CTA",
    "Insights blog system with article index (tag filtering, search, load-more "
    "pagination) and individual article pages (prose typography, social sharing, "
    "related articles)",
    "WordPress CMS integration \u2014 blog content managed via "
    "cms.playbook-group.co.uk with ISR refreshing content every 60 seconds",
    "Contact form on homepage and dedicated /contact page, both submitting to a "
    "custom WordPress REST endpoint with email notifications and admin-visible "
    "submission log",
    "Contact page with form, office details, and embedded Google My Maps showing "
    "Spencer Yard office, railway station, and Bath Street car park",
    "Site-wide navigation with scroll state, mobile drawer, focus trap, and "
    "progress bar",
    "Site-wide footer with navigation columns and legal strip",
    "Privacy policy and terms and conditions pages",
    "Session-once branded splash screen",
    "Full design token system mapped to Tailwind CSS v4",
    "Deployment pipeline \u2014 GitHub to Vercel with automatic builds",
]
for item in delivered_items:
    add_bullet(item)

add_para(
    "The following items were agreed as out of scope for Phase 1 (deferred):",
    space_before=Pt(12),
)
deferred_items = [
    "Service pages (/services/[slug]) \u2014 five pages, currently placeholder stubs",
    "The Playbook Model page (/the-playbook-model) \u2014 currently a placeholder stub",
    "Case studies section",
    "About page",
    "Google Analytics setup",
    "Staging/preview environment",
]
for item in deferred_items:
    add_bullet(item)


# =========================================================================
# 3. TECHNOLOGY STACK
# =========================================================================
doc.add_heading("Technology stack", level=1)
add_para(
    "The site is built with a modern, industry-standard technology stack chosen "
    "for performance, maintainability, and ease of content management."
)
make_table(
    ["Layer", "Technology", "Version / detail"],
    [
        ["Framework", "Next.js (App Router)", "16.1.7"],
        ["UI library", "React", "19.2.3"],
        ["Language", "TypeScript", "Strict mode"],
        ["Styling", "Tailwind CSS", "v4 (tokens via @theme directive)"],
        ["CMS", "WordPress", "Self-hosted at cms.playbook-group.co.uk"],
        ["Frontend hosting", "Vercel", "Automatic deploys from GitHub"],
        ["CMS hosting", "GoDaddy Managed WordPress", "cms.playbook-group.co.uk"],
        ["Maps", "Google My Maps (embed)", "Contact page"],
        ["Font", "Inter", "Google Fonts via next/font"],
        ["Version control", "GitHub", "github.com/stuart-byte-pm/playbook"],
    ],
    col_widths=[4.5, 5.5, 7.0],
)


# =========================================================================
# 4. ARCHITECTURE OVERVIEW
# =========================================================================
doc.add_heading("Architecture overview", level=1)
add_para(
    "The website uses a \u201cheadless\u201d architecture, which means the frontend "
    "(what visitors see) and the content management system (where you write blog "
    "posts) are separate systems that talk to each other."
)

doc.add_heading("How it fits together", level=2)
add_para(
    "The Next.js frontend application is hosted on Vercel. When code changes are "
    "pushed to the main branch on GitHub, Vercel automatically rebuilds and deploys "
    "the site. Most pages (homepage, sectors, contact) are pre-built at deploy time "
    "for maximum speed \u2014 this is called static site generation (SSG)."
)
add_para(
    "Blog content (the /insights section) is managed in WordPress at "
    "cms.playbook-group.co.uk. The frontend fetches articles from WordPress via its "
    "REST API. To keep the site fast while also showing recent content, the insights "
    "pages use incremental static regeneration (ISR) \u2014 the site checks WordPress "
    "for updated content every 60 seconds and rebuilds those pages automatically. "
    "You do not need to trigger a manual deploy when publishing a new blog post."
)
add_para(
    "Contact form submissions from both the homepage and the /contact page are sent "
    "to a custom WordPress REST endpoint. WordPress stores every submission as an "
    "internal record (visible under \u201cForm submissions\u201d in the WordPress "
    "admin dashboard) and sends an email notification to hello@playbook-group.co.uk."
)

doc.add_heading("Key architectural decisions", level=2)
for bp, rest in [
    ("Headless WordPress: ",
     "Gives full design control on the frontend while keeping WordPress as a "
     "familiar content editing tool."),
    ("ISR at 60 seconds: ",
     "Blog changes appear on the live site within a minute, without needing a full "
     "redeploy."),
    ("Server Actions for forms: ",
     "Form submissions are handled securely on the server, never exposed to the "
     "browser."),
    ("No database on the frontend: ",
     "All persistent data (blog posts, form submissions) lives in WordPress. The "
     "frontend is stateless."),
]:
    add_bullet(rest, bold_prefix=bp)


# =========================================================================
# 5. HOSTING & INFRASTRUCTURE
# =========================================================================
doc.add_heading("Hosting & infrastructure", level=1)

doc.add_heading("Frontend \u2014 Vercel", level=2)
add_para(
    "The website frontend is hosted on Vercel (vercel.com), a platform purpose-built "
    "for Next.js applications. Vercel handles SSL certificates, global CDN "
    "distribution, and automatic deployments from GitHub."
)
make_table(
    ["Detail", "Value"],
    [
        ["Provider", "Vercel"],
        ["Plan / tier", "[TO BE CONFIRMED]"],
        ["Account owner", "[TO BE CONFIRMED]"],
        ["Deployment trigger", "Automatic on push to main branch on GitHub"],
        ["Region", "Vercel global edge network"],
        ["SSL", "Automatic (managed by Vercel)"],
    ],
    col_widths=[5.0, 12.0],
)

doc.add_heading("CMS & form backend \u2014 GoDaddy Managed WordPress", level=2)
add_para(
    "WordPress is hosted on GoDaddy Managed WordPress. This is where blog content "
    "is created and where contact form submissions are stored."
)
make_table(
    ["Detail", "Value"],
    [
        ["Provider", "GoDaddy Managed WordPress"],
        ["Plan / tier", "[TO BE CONFIRMED]"],
        ["URL", "cms.playbook-group.co.uk"],
        ["Admin URL", "cms.playbook-group.co.uk/wp-admin"],
        ["SFTP host", "f05.c16.myftpupload.com"],
        ["Account owner", "[TO BE CONFIRMED]"],
    ],
    col_widths=[5.0, 12.0],
)

doc.add_heading("Source code \u2014 GitHub", level=2)
add_para("All source code is stored in a private GitHub repository.")
make_table(
    ["Detail", "Value"],
    [
        ["Repository", "github.com/stuart-byte-pm/playbook"],
        ["Branch strategy", "Single main branch \u2014 push to main triggers deploy"],
        ["Access", "[TO BE CONFIRMED]"],
    ],
    col_widths=[5.0, 12.0],
)


# =========================================================================
# 6. DEPLOYMENT & CI/CD
# =========================================================================
doc.add_heading("Deployment & CI/CD", level=1)

doc.add_heading("Frontend deployment", level=2)
add_para("The frontend is deployed automatically via Vercel:")
for s in [
    "A code change is pushed to the main branch on GitHub",
    "Vercel detects the push and starts a new build (typically 1\u20132 minutes)",
    "If the build succeeds, the new version goes live immediately",
    "If the build fails, the previous version remains live \u2014 nothing breaks",
]:
    add_bullet(s)
add_para(
    "There is no staging environment at present. All changes go directly to "
    "production. We recommend adding a staging/preview setup as a Phase 2 "
    "improvement.",
    space_before=Pt(8),
)

doc.add_heading("WordPress plugin deployment", level=2)
add_para(
    "The custom form-handling plugin (playbook-form-handler.php) is deployed to "
    "WordPress via SFTP. If this file ever needs updating:"
)
for s in [
    "Connect to f05.c16.myftpupload.com via an SFTP client (e.g. FileZilla)",
    "Navigate to wp-content/mu-plugins/",
    "Upload the updated playbook-form-handler.php file",
    "The change takes effect immediately \u2014 no restart needed",
]:
    add_bullet(s)
add_para("Contact Byte if you need changes to this plugin.", space_before=Pt(8))


# =========================================================================
# 7. DOMAIN & DNS
# =========================================================================
doc.add_heading("Domain & DNS", level=1)
make_table(
    ["Detail", "Value"],
    [
        ["Primary domain", "playbook-group.co.uk"],
        ["CMS subdomain", "cms.playbook-group.co.uk"],
        ["Registrar", "[TO BE CONFIRMED]"],
        ["Nameservers", "[TO BE CONFIRMED]"],
        ["SSL certificate (frontend)", "Managed automatically by Vercel"],
        ["SSL certificate (CMS)", "[TO BE CONFIRMED \u2014 likely managed by GoDaddy]"],
    ],
    col_widths=[5.0, 12.0],
)
add_placeholder(
    "[TO BE CONFIRMED: Please provide the domain registrar and nameserver details "
    "so this section can be completed.]"
)


# =========================================================================
# 8. SIGN-IN CREDENTIALS & ACCESS
# =========================================================================
doc.add_heading("Sign-in credentials & access", level=1)
add_para(
    "Below is every system your team may need access to. For security, passwords "
    "are never included in this document."
)
make_table(
    ["System", "URL", "How to access"],
    [
        ["Vercel dashboard", "vercel.com",
         "[TO BE CONFIRMED \u2014 ask Byte for login details]"],
        ["GitHub repository", "github.com/stuart-byte-pm/playbook",
         "[TO BE CONFIRMED \u2014 ask Byte for access]"],
        ["WordPress admin", "cms.playbook-group.co.uk/wp-admin",
         "[TO BE CONFIRMED \u2014 ask Byte for login details]"],
        ["GoDaddy hosting panel", "godaddy.com",
         "[TO BE CONFIRMED \u2014 ask Byte for login details]"],
        ["Domain registrar", "[TO BE CONFIRMED]", "[TO BE CONFIRMED]"],
        ["SFTP (WordPress)", "f05.c16.myftpupload.com",
         "[TO BE CONFIRMED \u2014 ask Byte for credentials]"],
    ],
    col_widths=[4.0, 6.0, 7.0],
)
add_placeholder(
    "[TO BE CONFIRMED: Please confirm where credentials are stored (e.g. 1Password "
    "vault, Bitwarden, shared secure note) so this can reference the correct "
    "location.]"
)


# =========================================================================
# 9. THIRD-PARTY INTEGRATIONS
# =========================================================================
doc.add_heading("Third-party integrations", level=1)

doc.add_heading("WordPress REST API", level=2)
add_para(
    "WordPress provides the content API for the insights blog. The frontend fetches "
    "published posts from the /wp-json/wp/v2/posts endpoint with embedded featured "
    "images and categories. This is a public, read-only API \u2014 no authentication "
    "is required for reading content."
)

doc.add_heading("WordPress custom REST endpoint (contact forms)", level=2)
add_para(
    "A custom endpoint at /wp-json/playbook/v1/contact receives form submissions "
    "from the website. This is handled by the playbook-form-handler.php mu-plugin. "
    "Submissions are stored as a custom post type in WordPress and trigger an email "
    "notification to hello@playbook-group.co.uk."
)

doc.add_heading("Google My Maps", level=2)
add_para(
    "The contact page embeds a Google My Maps iframe showing the Spencer Yard office "
    "location, Leamington Spa railway station, and Bath Street car park. This is a "
    "free embed with no API key required."
)
add_para(
    "Note: Google My Maps embeds require third-party cookies. The map will not "
    "display in incognito/private browsing windows or if third-party cookies are "
    "blocked.",
    color=GREY, size=Pt(10),
)

doc.add_heading("Google Analytics", level=2)
add_para(
    "Google Analytics has been deferred and is not yet configured. A placeholder "
    "comment exists in the root layout file ready for the Google Analytics script "
    "to be added. See the \u201cNext steps\u201d section for recommendations."
)

doc.add_heading("Sanity Studio (legacy \u2014 not in use)", level=2)
add_para(
    "The codebase includes Sanity CMS packages and an embedded Studio at /studio. "
    "This was part of the original architecture but was replaced by WordPress before "
    "content integration. The Sanity packages remain in devDependencies and the "
    "/studio route still exists, but neither is actively used. These can be safely "
    "removed in a future cleanup."
)


# =========================================================================
# 10. CONTENT MANAGEMENT
# =========================================================================
doc.add_heading("Content management", level=1)
add_para(
    "Blog content is managed entirely through WordPress. You do not need any "
    "technical knowledge to create or edit articles."
)

doc.add_heading("Logging in", level=2)
add_para(
    "Go to cms.playbook-group.co.uk/wp-admin and sign in with your WordPress "
    "credentials."
)

doc.add_heading("Creating a new blog post", level=2)
for s in [
    "In the WordPress dashboard, go to Posts \u2192 Add New",
    "Write your article using the WordPress block editor",
    "Assign a category (Governance, Healthcare, Regeneration, or Capital "
    "programmes) \u2014 this determines which tag filter it appears under on the "
    "website",
    "Upload a featured image \u2014 this appears as the cover image on the article "
    "card and article page",
    "Click Publish",
    "The article will appear on the website within approximately 60 seconds",
]:
    add_bullet(s)

doc.add_heading("Making an article \u201cfeatured\u201d", level=2)
add_para(
    "To pin an article as the featured post on the insights landing page, mark it "
    "as a \u201cSticky\u201d post in WordPress (found in the Post settings panel "
    "under \u201cStick to the top of the blog\u201d). Only one post should be "
    "sticky at a time."
)

doc.add_heading("Categories used on the site", level=2)
make_table(
    ["WordPress category", "Website tag"],
    [
        ["Governance", "Governance"],
        ["Healthcare", "Healthcare"],
        ["Regeneration", "Regeneration"],
        ["Capital programmes", "Capital programmes"],
    ],
    col_widths=[8.5, 8.5],
)

doc.add_heading("Viewing form submissions", level=2)
add_para(
    "All contact form submissions are visible in the WordPress admin under "
    "\u201cForm submissions\u201d in the left-hand menu. Each entry shows the "
    "sender\u2019s name, email, organisation, message, and which form it came from "
    "(homepage or contact page)."
)


# =========================================================================
# 11. MAKING CHANGES & UPDATES
# =========================================================================
doc.add_heading("Making changes & updates", level=1)

doc.add_heading("Content changes (self-serve)", level=2)
add_para(
    "You can update blog content, publish new articles, and manage categories "
    "directly in WordPress. Changes appear on the live site within approximately "
    "60 seconds."
)

doc.add_heading("Design or code changes", level=2)
add_para(
    "Any changes to the website design, layout, functionality, or page structure "
    "require code changes. Please contact Byte to discuss and scope these \u2014 "
    "our details are at the end of this document."
)

doc.add_heading("Adding new pages", level=2)
add_para(
    "New pages (e.g. service pages, case studies) require development work in the "
    "Next.js codebase. WordPress is used for blog articles only, not for creating "
    "new site pages."
)


# =========================================================================
# 12. ONGOING MAINTENANCE
# =========================================================================
doc.add_heading("Ongoing maintenance", level=1)
add_para("To keep the site secure and running smoothly, we recommend the following:")
make_table(
    ["Item", "Frequency", "Who"],
    [
        ["Vercel hosting renewal", "Per billing cycle", "[TO BE CONFIRMED]"],
        ["GoDaddy WordPress hosting renewal", "Per billing cycle", "[TO BE CONFIRMED]"],
        ["Domain registration renewal", "Annually", "[TO BE CONFIRMED]"],
        ["SSL certificate renewal",
         "Automatic (Vercel) / check GoDaddy", "Automatic / [TO BE CONFIRMED]"],
        ["WordPress core updates", "Monthly (check for updates)", "Byte or client"],
        ["WordPress plugin updates", "Monthly", "Byte or client"],
        ["Next.js / dependency updates", "Quarterly", "Byte"],
        ["Security review", "Quarterly", "Byte (recommended)"],
        ["Browser/device testing", "After any significant update", "Byte"],
    ],
    col_widths=[6.0, 4.5, 6.5],
)
add_placeholder(
    "[TO BE CONFIRMED: Ongoing support arrangement to be agreed \u2014 contact Byte "
    "to discuss options.]"
)


# =========================================================================
# 13. KNOWN ISSUES & LIMITATIONS
# =========================================================================
doc.add_heading("Known issues & limitations", level=1)
add_para(
    "We believe in being upfront about where things stand. The following items "
    "are known:"
)
known_issues = [
    ("Phone number unconfirmed: ",
     "The contact page currently shows \u201cNumber to be confirmed\u201d. Update "
     "this in the codebase once a phone number is available."),
    ("wp_mail() delivery: ",
     "Email notifications for form submissions use WordPress\u2019s built-in "
     "wp_mail() function. We have encountered some delivery reliability issues. If "
     "these persist, we recommend switching to Resend (a dedicated email delivery "
     "service) as an alternative. Byte can implement this."),
    ("Google Analytics not configured: ",
     "Google Analytics has been deferred. The site currently has no visitor "
     "tracking. A placeholder is ready in the code for when this is set up."),
    ("Service pages are stubs: ",
     "The five service pages at /services/[slug] are currently placeholder pages. "
     "These need content and design work to complete."),
    ("Playbook Model page is a stub: ",
     "The /the-playbook-model page is a placeholder. This needs a dedicated design "
     "and build."),
    ("Google My Maps in incognito: ",
     "The embedded map on the contact page requires third-party cookies. It will "
     "show a \u201crefused to connect\u201d message in incognito/private browsing. "
     "This is a Google limitation, not a site bug."),
    ("No staging environment: ",
     "All code changes deploy directly to production. We recommend adding a "
     "preview/staging setup for safer testing."),
    ("Legacy Sanity packages: ",
     "Sanity CMS packages remain in devDependencies from the original architecture "
     "plan. They are not used and can be removed in a future cleanup to reduce the "
     "dependency footprint."),
]
for bp, rest in known_issues:
    add_bullet(rest, bold_prefix=bp)


# =========================================================================
# 14. NEXT STEPS & RECOMMENDATIONS
# =========================================================================
doc.add_heading("Next steps & recommendations", level=1)
add_para(
    "No additional scope has been formally discussed for Phase 2 yet. Below are our "
    "recommendations for what to consider next, in rough priority order:"
)
recommendations = [
    ("Google Analytics setup: ",
     "Industry-standard analytics for traffic and conversion measurement. Free to "
     "use. Note: Google Analytics uses cookies, so a cookie consent banner will "
     "need to be added to the site before going live. We recommend setting this "
     "up before any marketing activity begins so you can measure results from "
     "day one."),
    ("Service pages build-out: ",
     "Complete the five service pages with full content, design, and imagery. These "
     "are the most visible gaps in the current site."),
    ("Playbook Model page: ",
     "A dedicated page for the five-stage Playbook Model \u2014 this is core to "
     "the brand proposition and should be built as a priority."),
    ("Email delivery (Resend): ",
     "If wp_mail() delivery issues persist, switching to Resend would give "
     "reliable, trackable email delivery for form submissions."),
    ("Spam protection: ",
     "Add honeypot fields or Cloudflare Turnstile to the contact forms to prevent "
     "bot submissions. We avoided reCAPTCHA as it would require a cookie consent "
     "banner."),
    ("Staging environment: ",
     "Set up a preview branch deployment on Vercel so code changes can be reviewed "
     "before going live."),
    ("Case studies section: ",
     "Once client engagements can be referenced, a case studies section would "
     "strengthen credibility."),
    ("About page: ",
     "A dedicated page introducing the team and advisory philosophy."),
    ("Performance audit: ",
     "A Lighthouse and Core Web Vitals audit to identify any optimisation "
     "opportunities before marketing campaigns."),
    ("Legacy cleanup: ",
     "Remove unused Sanity packages and the /studio route from the codebase."),
]
for bp, rest in recommendations:
    add_bullet(rest, bold_prefix=bp)


# =========================================================================
# 15. CONTACT & SUPPORT
# =========================================================================
doc.add_heading("Contact & support", level=1)
add_para(
    "If you need help with anything covered in this document, or want to discuss "
    "future work, please get in touch with the Byte team:"
)
make_table(
    ["", ""],
    [
        ["Agency", "Byte"],
        ["Website", "byte-pm.com"],
        ["Email", "stuart@byte-pm.com"],
        ["Primary contact", "Stuart McGreavy"],
    ],
    col_widths=[5.0, 12.0],
)
add_placeholder(
    "[TO BE CONFIRMED: Response time expectations and support arrangement to be "
    "agreed.]"
)
add_para(
    "For urgent issues (site down, broken functionality), email us with "
    "\u201cURGENT\u201d in the subject line and we will respond as quickly as "
    "possible.",
    space_before=Pt(12),
)
add_para(
    "For content updates (new blog posts, category changes, featured article), "
    "you can make these changes yourself via WordPress at "
    "cms.playbook-group.co.uk/wp-admin.",
    space_before=Pt(6),
)


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
for i, section in enumerate(doc.sections):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    if i > 0:
        run_left = fp.add_run("Byte  |  byte-pm.com")
        run_left.font.name = "Arial"
        run_left.font.size = Pt(8)
        run_left.font.color.rgb = GREY


# ---------------------------------------------------------------------------
# Save & validate
# ---------------------------------------------------------------------------
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(script_dir)
output_path = os.path.join(
    project_root, "assets", "Playbook_Advisory_Group_Handover_2026-04.docx"
)

doc.save(output_path)
print(f"Saved: {output_path}")

# Quick validation
doc2 = Document(output_path)
print(f"Paragraphs: {len(doc2.paragraphs)}")
print(f"Tables: {len(doc2.tables)}")
print(f"Sections: {len(doc2.sections)}")
print("OK")
